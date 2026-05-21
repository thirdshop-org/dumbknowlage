from __future__ import annotations

from pathlib import Path


def read_file(path: Path) -> str:
    ext = path.suffix.lower()
    readers = {
        ".txt": _read_text,
        ".md": _read_text,
        ".csv": _read_text,
        ".json": _read_text,
        ".xml": _read_text,
        ".html": _read_text,
        ".htm": _read_text,
        ".pdf": _read_pdf,
        ".docx": _read_docx,
    }
    reader = readers.get(ext)
    if reader is None:
        raise ValueError(f"Format non supporté: {ext} (supporté: {', '.join(readers)})")
    return reader(path)


def get_metadata(path: Path, text: str) -> dict:
    ext = path.suffix.lower()
    title = path.stem
    author = ""
    pages = 0

    if ext == ".pdf":
        import fitz
        doc = fitz.open(str(path))
        meta = doc.metadata
        title = meta.get("title", "") or path.stem
        author = meta.get("author", "")
        pages = doc.page_count
        doc.close()
    elif ext == ".docx":
        from docx import Document
        doc = Document(str(path))
        props = doc.core_properties
        title = props.title or path.stem
        author = props.author or ""
        pages = 0

    word_count = len(text.split())
    return {
        "title": title,
        "author": author,
        "pages": pages,
        "word_count": word_count,
        "file_type": ext.lstrip(".").upper(),
    }


def chunk_text(text: str, max_chars: int = 2000) -> list[dict]:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[dict] = []
    current = ""
    idx = 0
    for para in paragraphs:
        if len(current) + len(para) > max_chars and current:
            chunks.append({
                "chunk_index": idx,
                "start_time": float(idx),
                "end_time": float(idx + 1),
                "text": current.strip(),
            })
            idx += 1
            current = para
        else:
            if current:
                current += "\n\n" + para
            else:
                current = para
    if current:
        chunks.append({
            "chunk_index": idx,
            "start_time": float(idx),
            "end_time": float(idx + 1),
            "text": current.strip(),
        })
    return chunks


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    import fitz
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)
